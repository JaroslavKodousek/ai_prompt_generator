from pathlib import Path
from typing import Optional
import pypdf
from PIL import Image
import pytesseract
from ..core.models import DocumentType

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class DocumentLoader:
    """Handles loading and preprocessing documents."""

    @staticmethod
    def detect_document_type(file_path: Path) -> DocumentType:
        """Detect document type from file extension."""
        suffix = file_path.suffix.lower()
        if suffix == '.pdf':
            return DocumentType.PDF
        elif suffix in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
            return DocumentType.IMAGE
        elif suffix in ['.txt', '.md']:
            return DocumentType.TEXT
        elif suffix == '.docx':
            return DocumentType.DOCX
        else:
            raise ValueError(f"Unsupported file type: {suffix}")

    @staticmethod
    def load_text_file(file_path: Path) -> str:
        """Load plain text file."""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

    @staticmethod
    def load_pdf(file_path: Path) -> str:
        """Extract text from PDF file."""
        text_parts = []
        with open(file_path, 'rb') as f:
            pdf_reader = pypdf.PdfReader(f)
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
        return "\n\n".join(text_parts)

    @staticmethod
    def load_image(file_path: Path) -> str:
        """Extract text from image using OCR."""
        image = Image.open(file_path)
        text = pytesseract.image_to_string(image)
        return text

    @staticmethod
    def load_docx(file_path: Path) -> str:
        """Extract text from DOCX file."""
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is not installed. Install it with: pip install python-docx"
            )

        doc = Document(file_path)
        text_parts = []

        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))

        return "\n".join(text_parts)

    @staticmethod
    def load(file_path: str | Path) -> tuple[str, DocumentType]:
        """
        Load document and return text content with document type.

        Args:
            file_path: Path to document file

        Returns:
            Tuple of (text_content, document_type)
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        doc_type = DocumentLoader.detect_document_type(file_path)

        if doc_type == DocumentType.PDF:
            text = DocumentLoader.load_pdf(file_path)
        elif doc_type == DocumentType.IMAGE:
            text = DocumentLoader.load_image(file_path)
        elif doc_type == DocumentType.TEXT:
            text = DocumentLoader.load_text_file(file_path)
        elif doc_type == DocumentType.DOCX:
            text = DocumentLoader.load_docx(file_path)
        else:
            raise ValueError(f"Unsupported document type: {doc_type}")

        return text, doc_type

    @staticmethod
    def preprocess_text(text: str, max_length: Optional[int] = None) -> str:
        """
        Preprocess text (clean up whitespace, truncate if needed).

        Args:
            text: Raw text content
            max_length: Optional maximum length to truncate to

        Returns:
            Preprocessed text
        """
        # Remove excessive whitespace
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(line for line in lines if line)

        # Truncate if needed
        if max_length and len(text) > max_length:
            text = text[:max_length] + "\n[... truncated]"

        return text
